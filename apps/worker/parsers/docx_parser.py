from docx import Document
from typing import Dict, Any
from io import BytesIO


class DOCXParser:
    """Extract text and metadata from DOCX files."""
    
    @staticmethod
    def extract_text(file_content: bytes) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        try:
            doc = Document(BytesIO(file_content))
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            full_text = "\n\n".join(paragraphs)
            
            # Extract tables
            tables_data = []
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_rows.append(row_data)
                tables_data.append(table_rows)
            
            # Get core properties
            core_props = doc.core_properties
            metadata = {
                "author": core_props.author,
                "title": core_props.title,
                "subject": core_props.subject,
                "created": str(core_props.created) if core_props.created else None,
                "modified": str(core_props.modified) if core_props.modified else None,
            }
            
            return {
                "success": True,
                "full_text": full_text,
                "paragraphs": paragraphs,
                "tables": tables_data,
                "metadata": metadata,
                "paragraph_count": len(paragraphs),
                "table_count": len(tables_data)
            }
        
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "full_text": "",
                "paragraphs": [],
                "tables": []
            }
